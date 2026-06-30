#!/usr/bin/env python3
"""Generate tailored DOCX resumes and cover letters for selected jobs."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = Path("generated_applications")
OUT.mkdir(exist_ok=True)

CANDIDATE = "Kundan Srinivas Sakkuru"
CONTACT = "Tampa, FL | +1 813 568 7378 | kundansrinivas377@gmail.com"
LINKS = "github.com/kundan198 | linkedin.com/in/kundan-srinivas-sakkuru | kundansrinivas.vercel.app"

SELECTED = [
    {
        "company": "Buyers Edge Platform",
        "title": "Junior Developer - Python & Go",
        "score": 8.8,
        "email": "recruiting@buyersedgeplatform.com",
        "manager": "Not provided by Apify",
        "reason": "Best 0-2 year fit. Strong Python, APIs, SQL, Docker, AWS, LangChain, OpenAI API, automation, and AI agent overlap.",
        "summary": "Entry-level software engineer with production Python, API, automation, cloud, and applied AI experience. Built real systems used in research and hackathon production settings, including serverless data pipelines, LLM-backed workflows, and full-stack dashboards. Strong fit for junior automation and AI agent work with Python, SQL, cloud services, and fast learning across new tools like Go and Palantir AIP.",
        "skills": [
            "Python, Java, TypeScript, JavaScript, SQL, Bash, Dart",
            "FastAPI, REST APIs, WebSocket, Firebase, PostgreSQL, MongoDB, Redis",
            "LangChain, RAG, ChromaDB, OpenAI API concepts, Hugging Face, Prompt Engineering",
            "Docker, AWS, GCP, CI/CD, GitHub Actions, unit testing, code review",
            "React, Next.js, Flutter, React Native",
        ],
        "cover": [
            "Buyers Edge Platform stands out because the Junior Developer role combines exactly the areas I have been building in: Python automation, data workflows, APIs, and practical AI agents that remove manual work. I am a Computer Science master's student at the University of South Florida with 1.5 years of internship and research engineering experience, and I am authorized to work in the United States without sponsorship.",
            "At SHIELD Lab, I built CogniX, a production Flutter and Firebase platform used by 20+ researchers, replacing manual study workflows and reducing data collection time by 80%. I also engineered Python Cloud Functions that automatically generated analysis-ready CSV and PDF reports, saving 3+ researcher hours per study. In BayShield, I built a FastAPI and WebSocket multi-agent disaster response system using LangChain, RAG, ChromaDB, and live data pipelines, which maps closely to your AI-first automation direction.",
            "I would be excited to bring my Python, API, SQL, cloud, and applied AI background to Buyers Edge Platform and grow quickly under senior engineering mentorship. I would welcome the chance to discuss how I can contribute to your automation and AI agent roadmap.",
        ],
        "email_body": {
            "subject": "Python automation fit",
            "body": "Hi [First Name],\nBuyers Edge Platform's Junior Developer role caught my attention because it combines Python automation, data workflows, and AI agents. I am a USF Computer Science master's student with 1.5 years of research engineering experience and production work across Python, Firebase, FastAPI, LangChain, and cloud workflows. At SHIELD Lab, I built automation that saved 3+ researcher hours per study and reduced data collection time by 80%. Would you be open to a 15 minute call or reviewing my application for this role? Thank you for your time, and I would be grateful to connect.",
        },
    },
    {
        "company": "UPMC",
        "title": "Software Engineer - Associate",
        "score": 8.2,
        "email": "recruiting@upmc.com",
        "manager": "Not provided by Apify",
        "reason": "Fully remote associate role with REST/Web API, SDLC, unit testing, Agile, documentation, and growth-oriented expectations.",
        "summary": "Entry-level software engineer focused on reliable application development, REST APIs, secure coding habits, unit testing, documentation, and collaborative delivery. Experienced building production research software and backend automation under real operational constraints. Strong communicator with Agile, code review, and cross-functional experience.",
        "skills": [
            "Python, Java, TypeScript, JavaScript, SQL, Bash, Dart",
            "REST APIs, FastAPI, Django, WebSocket, Firebase, PostgreSQL, MongoDB",
            "Azure-ready cloud fundamentals, AWS, GCP, Docker, CI/CD, GitHub Actions",
            "Agile/Scrum, SDLC, unit testing, technical documentation, code review",
            "React, Next.js, Flutter, secure and maintainable application development",
        ],
        "cover": [
            "UPMC's Software Engineer - Associate role is appealing because it is built for a developer who can learn quickly, contribute clean code, document work clearly, and grow within a disciplined application engineering team. My background fits that path well: I am a USF Computer Science master's student with production research software experience, strong REST API fundamentals, and authorization to work in the United States without sponsorship.",
            "At SHIELD Lab, I built and deployed CogniX, a Flutter and Firebase platform used in live neuroscience studies by 20+ researchers, with offline-first sync and zero data loss across 50+ production sessions. I also created Python Cloud Functions that delivered CSV and PDF outputs automatically after each session, saving 3+ researcher hours per study. Across BayShield and SignBridge, I practiced SDLC basics through backend APIs, real-time systems, testing, documentation, and cross-functional delivery.",
            "I would be glad to bring my application development discipline, API experience, and growth mindset to UPMC's remote engineering team. I would welcome an interview to discuss how I can contribute as an associate engineer.",
        ],
        "email_body": {
            "subject": "Associate engineer fit",
            "body": "Hi [First Name],\nUPMC's Software Engineer - Associate role stood out because it emphasizes REST APIs, SDLC fundamentals, documentation, and growth within a remote team. I am a USF Computer Science master's student with 1.5 years of research engineering experience building production software used in live studies. At SHIELD Lab, I delivered a Flutter and Firebase platform with offline sync and Python reporting automation that saved 3+ researcher hours per study. Would you be open to a 15 minute call or reviewing my application? Thank you, and I appreciate your time.",
        },
    },
    {
        "company": "BeaconFire Inc.",
        "title": "Software Engineer",
        "score": 8.0,
        "email": "recruiting@beaconfireinc.com",
        "manager": "Not provided by Apify",
        "reason": "Explicit 0-1 year fit with Java, JavaScript, React, SQL, Spring/RAG/AI concepts, and a master's degree requirement.",
        "summary": "Software engineer with master's-level CS training, production full-stack experience, and strong Java, JavaScript, React, SQL, API, and AI/RAG fundamentals. Built reliable applications and AI systems using FastAPI, LangChain, ChromaDB, Firebase, PostgreSQL, and React. Ready for an entry-level Java and AI integration role with strong learning velocity.",
        "skills": [
            "Java, Python, TypeScript, JavaScript, SQL, Dart",
            "React, Next.js, REST APIs, FastAPI, Django, WebSocket",
            "PostgreSQL, MongoDB, Firebase, Docker, CI/CD, GitHub Actions",
            "LangChain, RAG, ChromaDB, Hugging Face, PyTorch, TensorFlow",
            "Agile/Scrum, OOP, data structures, unit testing, code review",
        ],
        "cover": [
            "BeaconFire's Software Engineer role is a strong match because it is designed for a 0-1 year engineer with Java, JavaScript, React, SQL, and AI integration interest. I am completing my MS in Computer Science at the University of South Florida and have 1.5 years of internship and research engineering experience building production software and AI systems.",
            "My strongest relevant work includes BayShield, a FastAPI and WebSocket disaster response platform with LangChain, RAG, ChromaDB, and a live React dashboard, and CogniX, a production Flutter and Firebase app used by 20+ researchers with Python Cloud Functions for automated reporting. These projects required OOP fundamentals, API design, database work, CI/CD habits, documentation, and practical AI integration under real constraints.",
            "I would be excited to contribute to BeaconFire's Java, full-stack, and AI integration work while continuing to grow as an entry-level software engineer. I would welcome the opportunity to interview and show how quickly I can contribute.",
        ],
        "email_body": {
            "subject": "Java AI engineer fit",
            "body": "Hi [First Name],\nBeaconFire's Software Engineer role caught my attention because it combines entry-level Java, React, SQL, and AI integration work. I am a USF Computer Science master's student with 1.5 years of research engineering experience and projects in FastAPI, React, LangChain, RAG, Firebase, and production automation. My BayShield project used LangChain and ChromaDB to ground multi-agent decisions in verified reports while serving a real-time React dashboard. Would you be open to a 15 minute call or reviewing my application? Thank you, and I would be glad to connect.",
        },
    },
    {
        "company": "Candid Health",
        "title": "Software Engineer",
        "score": 7.8,
        "email": "recruiting@candidhealth.com",
        "manager": "Not provided by Apify",
        "reason": "High-quality YC company with strong Python, PostgreSQL, Docker, React, TypeScript, GCP, customer problem solving, and ML/data overlap.",
        "summary": "Full-stack software engineer with strong Python, TypeScript, React, PostgreSQL, Docker, cloud, and applied AI experience. Motivated by operationally complex domains where clean software can remove manual work. Built production research workflows, backend automation, real-time systems, and AI-assisted platforms with measurable outcomes.",
        "skills": [
            "Python, TypeScript, JavaScript, Java, SQL, Dart",
            "React, Next.js, FastAPI, REST APIs, WebSocket, Django",
            "PostgreSQL, Firebase, MongoDB, Redis, Docker, GCP, AWS",
            "PyTorch, TensorFlow, LangChain, RAG, ChromaDB, Prompt Engineering",
            "System design, Agile/Scrum, unit testing, code review, technical writing",
        ],
        "cover": [
            "Candid Health's mission resonates with me because medical billing is exactly the kind of complex, high-friction workflow where strong software can create real operational leverage. I am a USF Computer Science master's student with 1.5 years of engineering experience, and my best work has been building systems that replace manual processes with reliable applications, automation, and data pipelines.",
            "At SHIELD Lab, I built CogniX, a production Flutter and Firebase platform used by 20+ researchers, replacing manual data collection workflows and reducing collection time by 80%. I also built Python Cloud Functions that generated CSV and PDF reports automatically, saving 3+ researcher hours per study. In BayShield, I built a FastAPI, WebSocket, LangChain, and ChromaDB system with a React dashboard, giving me hands-on experience across backend services, real-time interfaces, data grounding, and AI-assisted workflows.",
            "I would be excited to bring my full-stack background, user-first mindset, and automation experience to Candid Health. I would welcome the opportunity to discuss how I can contribute to your engineering team.",
        ],
        "email_body": {
            "subject": "Candid full-stack fit",
            "body": "Hi [First Name],\nCandid Health's software engineering role stood out because you are using software and data to simplify one of healthcare's most complex operational problems. I am a USF Computer Science master's student with production experience across Python, TypeScript, React, Firebase, FastAPI, and applied AI workflows. At SHIELD Lab, I built a platform that reduced research data collection time by 80% and automated reporting with Python Cloud Functions. Would you be open to a 15 minute call or reviewing my application? Thank you, and I would be excited to connect.",
        },
    },
    {
        "company": "Reserv",
        "title": "Full Stack Engineer",
        "score": 7.1,
        "email": "recruiting@reserv.com",
        "manager": "Not provided by Apify",
        "reason": "Remote-friendly full-stack role with React, TypeScript, backend ownership, AI/automation domain, and product engineering overlap, with Ruby as the main gap.",
        "summary": "Full-stack engineer with production experience across React, TypeScript, backend APIs, Firebase, Python automation, and applied AI. Comfortable shipping end-to-end product features, writing precise technical documentation, and collaborating across product, research, and engineering stakeholders. Strong fit for early-stage product engineering with fast learning in new backend stacks.",
        "skills": [
            "TypeScript, JavaScript, Python, Java, SQL, Dart",
            "React, Next.js, FastAPI, REST APIs, WebSocket, Django",
            "Firebase, PostgreSQL, MongoDB, Docker, AWS, GCP, CI/CD",
            "LangChain, RAG, ChromaDB, PyTorch, TensorFlow, Prompt Engineering",
            "Product engineering, Agile/Scrum, technical writing, code review",
        ],
        "cover": [
            "Reserv's Full Stack Engineer role is exciting because it combines product engineering, automation, AI, and end-to-end ownership in a domain where better software can change the claims experience. I have built production software in research settings and AI-driven hackathon systems, and I enjoy turning ambiguous workflows into reliable tools that real users can use.",
            "At SHIELD Lab, I shipped CogniX, a cross-platform Flutter and Firebase app used by 20+ researchers in live neuroscience studies, with offline-first sync and zero data loss across 50+ production sessions. I also automated reporting with Python Cloud Functions, saving 3+ researcher hours per study. My BayShield project added full-stack AI experience through FastAPI, WebSocket, LangChain, ChromaDB, and a real-time React dashboard.",
            "Although my strongest backend production experience is Python rather than Ruby, I have repeatedly learned new stacks quickly and shipped under real constraints. I would welcome the chance to discuss how I can contribute to Reserv's full-stack product work.",
        ],
        "email_body": {
            "subject": "Full-stack automation fit",
            "body": "Hi [First Name],\nReserv's Full Stack Engineer role caught my attention because it blends product engineering, automation, and AI in insurance claims. I am a USF Computer Science master's student with production experience across React, TypeScript, Firebase, Python automation, FastAPI, and real-time systems. At SHIELD Lab, I shipped a platform used in live studies with offline sync, zero data loss across 50+ sessions, and automated reporting that saved 3+ researcher hours per study. Would you be open to a 15 minute call or reviewing my application? Thank you, and I would be glad to connect.",
        },
    },
]


def clean(text: str) -> str:
    return text.replace("—", "-").replace("–", "-").replace("‑", "-").replace("’", "'").replace("“", '"').replace("”", '"')


def slug(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")


def setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.03
    for name, size in [("Heading 1", 12.5), ("Heading 2", 11.2)]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(31, 77, 120)
        st.paragraph_format.space_before = Pt(6)
        st.paragraph_format.space_after = Pt(2)


def add_header(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CANDIDATE.upper())
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(11, 37, 69)
    p.paragraph_format.space_after = Pt(1)
    for line in [CONTACT, LINKS, subtitle]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(clean(line))
        r.font.size = Pt(9)
        if line == subtitle:
            r.italic = True
        p.paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(clean(item))


def resume_doc(job: dict) -> Path:
    doc = Document()
    setup(doc)
    add_header(doc, f"Tailored for {job['title']} at {job['company']}")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(clean(job["summary"]))
    doc.add_heading("Technical Skills", level=1)
    add_bullets(doc, job["skills"])
    doc.add_heading("Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Software Engineer - Research | SHIELD Lab, University of South Florida | Tampa, FL | Aug 2025 - May 2026").bold = True
    add_bullets(doc, [
        "Architected and deployed CogniX, a production cross-platform Flutter app used by 20+ researchers across live neuroscience studies, replacing manual processes and reducing data collection time by 80%.",
        "Designed a zero-collision Firestore schema with offline-first persistence and automatic sync, achieving zero data loss across 50+ production sessions in low-connectivity field environments.",
        "Built a serverless Python Cloud Function pipeline that auto-delivered analysis-ready CSV and PDF reports on session completion, saving 3+ researcher hours per study.",
        "Partnered with research stakeholders to translate ambiguous workflow requirements into tested product behavior, technical documentation, and reliable release updates.",
    ])
    doc.add_heading("Selected Projects", level=1)
    projects = [
        ("BayShield - Agentic AI Disaster Response System | Python, LangChain, RAG, ChromaDB, FastAPI, WebSocket, React | HackUSF 2026, 3rd Place", [
            "Designed an event-driven FastAPI and WebSocket architecture for real-time multi-agent updates using live Hurricane Helene data.",
            "Built a RAG pipeline with ChromaDB and sentence-transformers to ground decisions in verified FEMA reports and reduce hallucination risk.",
            "Delivered a React dashboard with 5-second refresh for incident synthesis, shelter optimization, and resource allocation workflows.",
        ]),
        ("SignBridge - Real-Time Generative AI Sign Language Platform | MediaPipe, Google Gemini, NLTK, Django, ElevenLabs | HackaBull 2025, 2nd Place", [
            "Built a two-way AI communication system for speech-to-sign and sign-to-speech with sub-second end-to-end latency.",
            "Integrated NLP grammar restructuring, landmark tracking, Gemini reasoning, and ElevenLabs TTS into one accessible workflow.",
        ]),
        ("Skin Cancer Prediction using CNN | Python, PyTorch, TensorFlow, ResNet50 | Published in Scopus", [
            "Trained a 7-class skin lesion classifier on HAM10000 and improved minority-class recall by 18% using inverse-frequency weighting.",
        ]),
    ]
    for title, bullets in projects:
        p = doc.add_paragraph()
        p.add_run(clean(title)).bold = True
        add_bullets(doc, bullets)
    doc.add_heading("Education", level=1)
    add_bullets(doc, [
        "MS, Computer Science, University of South Florida, Tampa, FL, Aug 2024 - May 2026. Coursework: Operating Systems, Algorithms, Machine Learning, Mobile Systems, Software Engineering.",
        "BE, Computer Science and Engineering, R.M.D Engineering College, Anna University, India, Nov 2020 - May 2024. CGPA: 9.03/10, Gold Medal, Best Outgoing Student.",
    ])
    doc.add_heading("Achievements", level=1)
    add_bullets(doc, [
        "HackaBull 2025, Standout Entry and 2nd Place for SignBridge.",
        "HackUSF 2026, 3rd Place for BayShield.",
        "Smart India Hackathon 2022 National Finalist, Cognizant PRODIGI Hackathon Finalist, TNSI Program Selected.",
    ])
    path = OUT / f"{slug(CANDIDATE)}_{slug(job['company'])}_Resume.docx"
    doc.save(path)
    return path


def cover_doc(job: dict) -> Path:
    doc = Document()
    setup(doc)
    add_header(doc, f"Cover Letter for {job['title']} at {job['company']}")
    doc.add_paragraph(clean(f"Dear {job['company']} Hiring Team,"))
    for para in job["cover"]:
        doc.add_paragraph(clean(para))
    doc.add_paragraph(clean("Sincerely,\nKundan Srinivas Sakkuru"))
    path = OUT / f"{slug(CANDIDATE)}_{slug(job['company'])}_Cover.docx"
    doc.save(path)
    return path


def main() -> None:
    manifest = []
    for job in SELECTED:
        resume = resume_doc(job)
        cover = cover_doc(job)
        manifest.append({**job, "resume_docx": str(resume), "cover_docx": str(cover)})
    with open(OUT / "manifest.json", "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
